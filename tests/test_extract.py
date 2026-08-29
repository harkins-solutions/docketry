import tempfile
import unittest
from pathlib import Path

from docketry.extract import Extraction, ExtractionError, Page, extract_path

try:
    import pypdf
    HAVE_PYPDF = True
except ImportError:
    HAVE_PYPDF = False
try:
    import docx
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


def minimal_pdf(texts):
    """Hand-assemble a valid single-font PDF, one page per text, correct xref."""
    objects = []
    kids = " ".join(f"{3 + 2 * i} 0 R" for i in range(len(texts)))
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(texts)} >>".encode())
    font_num = 3 + 2 * len(texts)
    for i, text in enumerate(texts):
        content_num = 4 + 2 * i
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_num} 0 R /Resources << /Font << /F1 {font_num} 0 R >> >> >>".encode()
        )
        stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"
        )
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for n, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{n} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n".encode()
    )
    return bytes(out)


class TestExtractText(unittest.TestCase):
    def test_txt(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "note.txt"
            p.write_text("hello world")
            r = extract_path(p)
            self.assertEqual(r.full_text, "hello world")
            self.assertEqual(r.method, "text")

    def test_unsupported_type_is_loud(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "img.png"
            p.write_bytes(b"x")
            with self.assertRaises(ExtractionError):
                extract_path(p)

    def test_missing_file_is_loud(self):
        with self.assertRaises(ExtractionError):
            extract_path("/nonexistent/file.txt")

    def test_bad_ocr_mode_refused(self):
        with self.assertRaises(ValueError):
            extract_path("x.txt", ocr="maybe")


@unittest.skipUnless(HAVE_PYPDF, "pypdf not installed")
class TestExtractPdf(unittest.TestCase):
    def test_two_page_pdf_with_page_map(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "brief.pdf"
            p.write_bytes(minimal_pdf(["First page words", "Second page words"]))
            r = extract_path(p, ocr="never")
            self.assertEqual(len(r.pages), 2)
            self.assertIn("First page words", r.pages[0].text)
            self.assertIn("Second page words", r.pages[1].text)
            self.assertEqual(r.pages[0].method, "native")
            offset = r.full_text.index("Second page")
            self.assertEqual(r.page_for_offset(offset), 2)
            self.assertEqual(r.page_for_offset(0), 1)

    def test_scanned_pdf_warns_when_ocr_is_not_used(self):
        """A page with no text and no OCR must say so, never return silence.

        This asked for ocr="auto" and relied on the machine having no
        tesseract, so it passed for a reason that had nothing to do with the
        behaviour under test — and started failing the moment CI installed the
        OCR toolchain. Ask for the no-OCR path explicitly instead.
        """
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "scan.pdf"
            p.write_bytes(minimal_pdf([""]))
            r = extract_path(p, ocr="never")
            self.assertTrue(any("no extractable text" in w for w in r.warnings))


@unittest.skipUnless(HAVE_DOCX, "python-docx not installed")
class TestExtractDocx(unittest.TestCase):
    def test_docx_paragraphs_and_tables(self):
        import docx as docx_mod
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "draft.docx"
            d = docx_mod.Document()
            d.add_paragraph("MOTION FOR SUMMARY JUDGMENT")
            d.add_paragraph("The undisputed facts follow.")
            table = d.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Exhibit A"
            table.rows[0].cells[1].text = "Policy"
            d.save(str(p))
            r = extract_path(p)
            self.assertIn("MOTION FOR SUMMARY JUDGMENT", r.full_text)
            self.assertIn("Exhibit A | Policy", r.full_text)
            self.assertTrue(any("no fixed pages" in w for w in r.warnings))


class TestPageForOffset(unittest.TestCase):
    def test_empty_extraction(self):
        self.assertIsNone(Extraction(pages=[], method="text").page_for_offset(0))


if __name__ == "__main__":
    unittest.main()


class TestCorruptFilesFailInOurVocabulary(unittest.TestCase):
    """A truncated PDF is an ordinary thing to receive, not a crash."""

    def test_a_truncated_pdf_raises_extraction_error_not_a_library_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "truncated.pdf"
            p.write_bytes(b"%PDF-1.4\n1 0 obj")     # header, then nothing
            with self.assertRaises(ExtractionError) as ctx:
                extract_path(p)
            self.assertIn("could not be read as a PDF", str(ctx.exception))

    def test_a_file_that_is_not_a_pdf_at_all(self):
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "notreally.pdf"
            p.write_bytes(b"x")
            with self.assertRaises(ExtractionError):
                extract_path(p)

    def test_a_file_that_is_not_a_docx_at_all(self):
        if not HAVE_DOCX:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as tmp:
            p = Path(tmp) / "notreally.docx"
            p.write_bytes(b"x")
            with self.assertRaises(ExtractionError) as ctx:
                extract_path(p)
            self.assertIn("Word document", str(ctx.exception))
