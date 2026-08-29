"""Timeline weave: layers stay apart, gaps stay honest, provenance survives."""
import json
import tempfile
import unittest
from pathlib import Path

from docketry.envelope import parse_message
from docketry.store import Store
from docketry.timeline import (
    CORRESPONDENCE,
    ATTACHED,
    CORRESPONDENCE,
    LINK_CAPTURED,
    PROVEN,
    RECORD,
    REFERENCED_ONLY,
    build,
    normalise_case_number,
)


def _raw(subject, body, *, msgid, frm="clerk@uscourts.gov", refs=None, reply=None):
    hdrs = [f"From: {frm}", "To: firm@example.com", f"Subject: {subject}",
            f"Message-ID: <{msgid}>", "Date: Mon, 3 Aug 2026 09:00:00 -0400"]
    if reply:
        hdrs.append(f"In-Reply-To: <{reply}>")
    if refs:
        hdrs.append("References: " + " ".join(f"<{r}>" for r in refs))
    return ("\r\n".join(hdrs) + "\r\n\r\n" + body).encode()


class TestNormalise(unittest.TestCase):
    def test_formatting_differences_collapse(self):
        self.assertEqual(normalise_case_number("8:26-cv-01234"),
                         normalise_case_number("826CV01234"))

    def test_different_numbers_stay_different(self):
        self.assertNotEqual(normalise_case_number("26-CA-001"),
                            normalise_case_number("26-CA-002"))


class TestBuild(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.store = Store(self.tmp)

    def _notice(self, msgid, fields, ntype="service_notice", adapter="pacer-nef",
                body="x", frm="clerk@uscourts.gov", refs=None):
        env = parse_message(_raw("Activity in Case", body, msgid=msgid, frm=frm,
                                 refs=refs),
                            source="intake", fetched_at="2026-08-03T09:00:00")
        mid = self.store.ingest(env, first_stage="ingest")
        self.store.add_notice(mid, adapter, ntype, fields, [])
        return mid

    def test_only_this_case_is_woven_in(self):
        self._notice("a@x", {"case_number": "8:26-cv-01234", "document_number": "1"})
        self._notice("b@x", {"case_number": "9:26-cv-99999", "document_number": "1"})
        tl = build(self.store, "826CV01234")
        self.assertEqual(len(tl.entries), 1)

    def test_every_entry_carries_its_source(self):
        mid = self._notice("c@x", {"case_number": "26-CA-1", "docket_text": "Order"})
        tl = build(self.store, "26-CA-1")
        self.assertEqual(tl.entries[0].source_message, mid)
        self.assertEqual(tl.entries[0].source_adapter, "pacer-nef")

    def test_availability_reflects_what_we_actually_hold(self):
        self._notice("d@x", {"case_number": "26-CA-2",
                             "document_link": "https://ecf.flsd.uscourts.gov/x"})
        self._notice("e@x", {"case_number": "26-CA-2", "docket_text": "Notice"})
        tl = build(self.store, "26-CA-2")
        kinds = {e.availability for e in tl.entries}
        self.assertIn(LINK_CAPTURED, kinds)
        self.assertIn(REFERENCED_ONLY, kinds)

    def test_a_hole_in_a_federal_sequence_is_a_proven_gap(self):
        for n in (45, 46, 48):
            self._notice(f"f{n}@x", {"case_number": "8:26-cv-1",
                                     "document_number": str(n)})
        tl = build(self.store, "8:26-cv-1")
        self.assertEqual(len(tl.gaps), 1)
        self.assertEqual(tl.gaps[0]["class"], PROVEN)
        self.assertEqual(tl.gaps[0]["numbers"], [47])

    def test_consecutive_missing_numbers_collapse_to_one_finding(self):
        # Served 1, 12 and 15: twelve missing numbers, but ONE thing to read.
        for n in (1, 12, 15):
            self._notice(f"r{n}@x", {"case_number": "8:26-cv-9",
                                     "document_number": str(n)})
        tl = build(self.store, "8:26-cv-9")
        self.assertEqual(len(tl.gaps), 1)
        self.assertEqual(tl.gaps[0]["count"], 12)
        self.assertIn("2-11", tl.gaps[0]["detail"])
        self.assertIn("13-14", tl.gaps[0]["detail"])

    def test_no_sequence_means_no_asserted_gap(self):
        # State notices carry no document number; silence beats a guess.
        for i in range(3):
            self._notice(f"g{i}@x", {"case_number": "26-CA-3",
                                     "documents": "Notice of Hearing"})
        tl = build(self.store, "26-CA-3")
        self.assertEqual(tl.gaps, [])

    def test_attached_threads_are_correspondence_never_record(self):
        self._notice("h@x", {"case_number": "26-CA-4", "docket_text": "Order"})
        env = parse_message(
            _raw("Re: settlement", "let's talk", msgid="t2@x",
                 frm="oc@otherfirm.com", refs=["root@x"]),
            source="intake", fetched_at="2026-08-03T10:00:00")
        self.store.ingest(env, first_stage="ingest")
        tl = build(self.store, "26-CA-4", threads=["root@x"])
        layers = {e.layer for e in tl.entries}
        self.assertEqual(layers, {RECORD, CORRESPONDENCE})
        corr = [e for e in tl.entries if e.layer == CORRESPONDENCE]
        self.assertEqual(len(corr), 1)
        self.assertFalse(corr[0].of_record)

    def test_undated_entries_sort_last_not_first(self):
        self._notice("i@x", {"case_number": "26-CA-5", "docket_text": "Dated"})
        tl = build(self.store, "26-CA-5")
        tl.entries.append(type(tl.entries[0])(
            when="", layer=RECORD, kind="service", title="Undated"))
        self.assertEqual(tl.sorted_entries()[-1].title, "Undated")


if __name__ == "__main__":
    unittest.main()


class TestExport(unittest.TestCase):
    def setUp(self):
        self.tmp = Path(tempfile.mkdtemp())
        self.store = Store(self.tmp)
        env = parse_message(
            _raw("Activity in Case", "x", msgid="x1@x"),
            source="intake", fetched_at="2026-08-03T09:00:00")
        mid = self.store.ingest(env, first_stage="ingest")
        self.store.add_notice(mid, "pacer-nef", "service_notice",
                              {"case_number": "8:26-cv-1", "docket_text": "Order",
                               "document_number": "3"}, [])
        self.tl = build(self.store, "8:26-cv-1")

    def test_xlsx_has_real_dates_filters_and_the_disclaimer(self):
        from openpyxl import load_workbook
        from docketry.export import to_xlsx
        out = to_xlsx(self.tl, self.tmp / "t.xlsx")
        wb = load_workbook(out)
        ws = wb["Timeline"]
        self.assertEqual(ws["A1"].value, "Date")
        self.assertEqual(ws.freeze_panes, "A2")
        self.assertIsNotNone(ws.auto_filter.ref)
        # A date stored as text cannot be sorted, which defeats a chronology.
        from datetime import datetime as _dt
        self.assertIsInstance(ws["A2"].value, _dt)
        self.assertIn("About this file", wb.sheetnames)
        self.assertIn("NOT the court's docket", wb["About this file"]["A3"].value)

    def test_docx_is_a_real_table_not_drawn_text(self):
        import docx
        from docketry.export import to_docx
        out = to_docx(self.tl, self.tmp / "t.docx")
        d = docx.Document(str(out))
        self.assertEqual(len(d.tables), 1)
        self.assertEqual(d.tables[0].rows[0].cells[0].text, "Date")
        self.assertEqual(len(d.tables[0].rows), 2)          # header + one entry
        body = "\n".join(p.text for p in d.paragraphs)
        self.assertIn("NOT the court's docket", body)

    def test_gaps_reach_both_exports(self):
        from openpyxl import load_workbook
        import docx
        from docketry.export import to_docx, to_xlsx
        self.tl.gaps = [{"class": "proven", "detail": "document 47 missing"}]
        wb = load_workbook(to_xlsx(self.tl, self.tmp / "g.xlsx"))
        about = "\n".join(str(c.value) for c in wb["About this file"]["A"])
        self.assertIn("document 47 missing", about)
        d = docx.Document(str(to_docx(self.tl, self.tmp / "g.docx")))
        self.assertIn("document 47 missing", "\n".join(p.text for p in d.paragraphs))


class TestReconcile(unittest.TestCase):
    def setUp(self):
        from docketry.timeline import Entry, Timeline
        self.tl = Timeline(case_number="8:26-cv-1")
        self.tl.entries = [
            Entry(when="2026-03-04T09:00:00", layer=RECORD, kind="service",
                  title="Complaint", doc_number=1),
            Entry(when="2026-03-19T09:00:00", layer=RECORD, kind="service",
                  title="Motion to Dismiss", doc_number=12),
            Entry(when="2026-03-21T09:00:00", layer=CORRESPONDENCE, kind="email",
                  title="Re: extension"),
        ]

    def test_csv_and_text_dockets_both_parse(self):
        from docketry.reconcile import parse_docket
        csv_rows = parse_docket("Doc,Date,Description\n1,03/04/2026,Complaint\n")
        txt_rows = parse_docket("  1  03/04/2026  Complaint\n  12 03/19/2026 Motion\n")
        self.assertEqual(csv_rows[0].doc_number, 1)
        self.assertEqual(len(txt_rows), 2)

    def test_unreadable_lines_are_skipped_not_guessed(self):
        from docketry.reconcile import parse_docket
        self.assertEqual(parse_docket("this is prose, not a docket"), [])

    def test_document_numbers_match_exactly(self):
        from docketry.reconcile import DocketLine, reconcile
        rec = reconcile(self.tl, [DocketLine(1, "03/04/2026", "Complaint"),
                                  DocketLine(12, "03/19/2026", "Mot. to Dismiss")])
        self.assertEqual(len(rec.matched), 2)
        self.assertTrue(rec.clean)

    def test_a_docket_entry_we_never_received_is_reported(self):
        from docketry.reconcile import DocketLine, reconcile
        rec = reconcile(self.tl, [DocketLine(1, "03/04/2026", "Complaint"),
                                  DocketLine(12, "03/19/2026", "MTD"),
                                  DocketLine(13, "03/20/2026", "Notice of Appearance")])
        self.assertEqual(len(rec.only_on_docket), 1)
        self.assertEqual(rec.only_on_docket[0].doc_number, 13)
        self.assertFalse(rec.clean)

    def test_something_we_hold_that_is_not_on_the_docket_is_reported(self):
        from docketry.reconcile import DocketLine, reconcile
        rec = reconcile(self.tl, [DocketLine(1, "03/04/2026", "Complaint")])
        self.assertEqual([e.doc_number for e in rec.only_here], [12])

    def test_correspondence_is_not_counted_as_a_discrepancy(self):
        from docketry.reconcile import DocketLine, reconcile
        rec = reconcile(self.tl, [DocketLine(1, "03/04/2026", "Complaint"),
                                  DocketLine(12, "03/19/2026", "MTD")])
        # The email was never going to be on a docket.
        self.assertEqual(rec.only_here, [])

    def test_fuzzy_matches_are_staged_never_merged(self):
        from docketry.timeline import Entry, Timeline
        from docketry.reconcile import DocketLine, reconcile
        tl = Timeline(case_number="26-CA-9")
        tl.entries = [Entry(when="2026-05-01T09:00:00", layer=RECORD,
                            kind="service", title="Notice of Hearing on Motion")]
        rec = reconcile(tl, [DocketLine(None, "05/01/2026",
                                        "Notice of Hearing on the Motion")])
        self.assertEqual(len(rec.to_confirm), 1)
        self.assertEqual(rec.matched, [])


class TestFindingsAndHelpers(unittest.TestCase):
    def test_a_rejected_filing_with_no_refile_is_surfaced(self):
        from docketry.timeline import Entry, Timeline, cross_layer_findings
        tl = Timeline(case_number="26-CA-7")
        tl.entries = [Entry(when="2026-05-01T09:00:00", layer=RECORD,
                            kind="filing", title="Filing Rejected — Envelope 9912")]
        findings = cross_layer_findings(tl)
        self.assertEqual(len(findings), 1)
        self.assertIn("REJECTED", findings[0])

    def test_a_rejected_filing_followed_by_a_refile_is_not_flagged(self):
        from docketry.timeline import Entry, Timeline, cross_layer_findings
        tl = Timeline(case_number="26-CA-8")
        tl.entries = [
            Entry(when="2026-05-01T09:00:00", layer=RECORD, kind="filing",
                  title="Filing Rejected — Envelope 9912"),
            Entry(when="2026-05-02T09:00:00", layer=RECORD, kind="filing",
                  title="Filing Accepted — Envelope 9931"),
        ]
        self.assertEqual(cross_layer_findings(tl), [])

    def test_threads_are_counted(self):
        from docketry.timeline import Entry, Timeline
        tl = Timeline(case_number="26-CA-10")
        tl.entries = [
            Entry(when="", layer=CORRESPONDENCE, kind="email", title="a",
                  thread_key="t1"),
            Entry(when="", layer=CORRESPONDENCE, kind="email", title="b",
                  thread_key="t1"),
            Entry(when="", layer=CORRESPONDENCE, kind="email", title="c",
                  thread_key="t2"),
        ]
        self.assertEqual(tl.threads(), {"t1": 2, "t2": 1})

    def test_an_unparseable_date_survives_as_text_not_a_blank(self):
        # A date we cannot parse must still reach the reader. Blanking it
        # would silently move an entry out of the chronology.
        from docketry.export import _as_date
        self.assertEqual(_as_date("filed sometime in March"),
                         "filed sometime in March")
        self.assertEqual(_as_date(""), "")

    def test_findings_reach_the_spreadsheet(self):
        from openpyxl import load_workbook
        from docketry.export import to_xlsx
        from docketry.timeline import Entry, Timeline
        tmp = Path(tempfile.mkdtemp())
        tl = Timeline(case_number="26-CA-11")
        tl.entries = [Entry(when="2026-05-01T09:00:00", layer=RECORD,
                            kind="filing", title="Something")]
        tl.findings = ["a filing was REJECTED and never re-filed"]
        wb = load_workbook(to_xlsx(tl, tmp / "f.xlsx"))
        about = "\n".join(str(c.value) for c in wb["About this file"]["A"])
        self.assertIn("never re-filed", about)


class TestClientMailReachesItsOwnLayer(unittest.TestCase):
    """The layer existed, was documented as privileged, and was always empty."""

    def setUp(self):
        self.store = Store(tempfile.mkdtemp())
        env = parse_message(
            _raw("Activity in Case", "x", msgid="n1@x"),
            source="t", fetched_at="now")
        mid = self.store.ingest(env, first_stage="ingest")
        self.store.add_notice(mid, "pacer-nef", "service_notice",
                              {"case_number": "26-CA-1"}, [])
        for msgid, frm in (("t1@x", "mr.doe@client.com"),
                           ("t2@x", "oc@theirfirm.com")):
            self.store.ingest(
                parse_message(_raw("Re: the case", "hello", msgid=msgid,
                                   frm=frm, refs=["root@x"]),
                              source="t", fetched_at="now"),
                first_stage="ingest")

    def _directory(self):
        from docketry.contacts import load_contacts
        p = Path(tempfile.mkdtemp()) / "c.toml"
        p.write_text('[[contact]]\nemail="mr.doe@client.com"\nkind="client"\n'
                     '[[contact]]\nemail="@theirfirm.com"\n'
                     'kind="opposing_counsel"\n')
        return load_contacts(p)

    def test_without_a_directory_everything_falls_to_correspondence(self):
        tl = build(self.store, "26-CA-1", threads=["root@x"])
        layers = {e.layer for e in tl.entries if e.kind != "service"}
        self.assertEqual(layers, {CORRESPONDENCE})

    def test_with_a_directory_client_mail_is_separated(self):
        from docketry.timeline import CLIENT
        tl = build(self.store, "26-CA-1", threads=["root@x"],
                   directory=self._directory())
        by_layer = {e.layer for e in tl.entries}
        self.assertIn(CLIENT, by_layer)
        client = [e for e in tl.entries if e.layer == CLIENT]
        self.assertEqual(len(client), 1)
        self.assertIn("client.com", client[0].actor)

    def test_the_other_side_stays_in_correspondence(self):
        tl = build(self.store, "26-CA-1", threads=["root@x"],
                   directory=self._directory())
        oc = [e for e in tl.entries if "theirfirm" in e.actor]
        self.assertEqual(oc[0].layer, CORRESPONDENCE)
        self.assertEqual(oc[0].kind, "opposing_counsel")

    def test_filtering_to_the_client_layer_returns_only_client_mail(self):
        from docketry.timeline import CLIENT
        tl = build(self.store, "26-CA-1", threads=["root@x"],
                   directory=self._directory())
        rows = tl.sorted_entries((CLIENT,))
        self.assertEqual(len(rows), 1)
