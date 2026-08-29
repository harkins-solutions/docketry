"""Timeline out to Word and Excel — the formats the work actually leaves in.

Two rules shape everything here.

Real tables, never drawn ones. A monospaced block that looks like a table in a
terminal is not a table in Word: it cannot be sorted, filtered, re-columned or
pasted into anything. Both exports produce native structures.

The disclaimer travels with the file. A reconstructed timeline exported to a
spreadsheet gets mailed around, printed, and eventually read by someone who
was not told what it is. So every export says on its own face that it was
assembled from what this firm received, that it is not the court's docket, and
that gaps are shown rather than filled — and every row carries the layer it
belongs to, so an inference can never be mistaken for a record.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from .timeline import DERIVED, Timeline

DISCLAIMER = (
    "Reconstructed from the notices, receipts and correspondence this firm "
    "received. This is NOT the court's docket and makes no claim to be "
    "complete. Gaps are listed as gaps, never filled in. Every row names the "
    "message it came from."
)

# (label, attribute, Excel character width, Word column width in inches).
# Word is the fussy one: it ignores column-level widths entirely, so the inch
# value has to be stamped onto every cell or the table "helpfully" autofits —
# which starves the Entry column, the only one anybody reads, and hands the
# space to Msg.
COLUMNS = [
    ("Date", "when", 20, 0.90),
    ("Layer", "layer", 16, 1.15),
    ("Type", "kind", 14, 0.70),
    ("Entry", "title", 62, 3.00),
    ("Party / From", "actor", 34, 1.65),
    ("Copy", "availability", 18, 0.80),
    ("Doc #", "doc_number", 9, 0.45),
    ("Source", "source_adapter", 18, 0.90),
    ("Msg", "source_message", 8, 0.45),
]

_AVAILABILITY_LABEL = {
    "attached": "held",
    "link_captured": "link only",
    "referenced_only": "not held",
}


def _cell(entry, attr):
    v = getattr(entry, attr, "")
    if attr == "availability":
        return _AVAILABILITY_LABEL.get(v, v)
    if attr == "layer":
        return str(v).capitalize()
    return "" if v is None else str(v)


def _as_date(iso: str):
    """Real date, or the raw string when it will not parse — never a silent blank."""
    if not iso:
        return ""
    try:
        return datetime.fromisoformat(iso.replace("Z", "+00:00")).replace(tzinfo=None)
    except ValueError:
        return iso


def to_xlsx(tl: Timeline, path: str | Path, *, layers=None, thread=None) -> Path:
    """A working spreadsheet: frozen header, filters on every column, real dates."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        raise RuntimeError(
            "Excel export needs the 'export' extra: pip install 'docketry[export]'"
        ) from None

    path = Path(path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Timeline"

    head = Font(bold=True, color="FFFFFF")
    fill = PatternFill("solid", fgColor="31425A")
    for i, (label, _attr, width, _in) in enumerate(COLUMNS, start=1):
        c = ws.cell(row=1, column=i, value=label)
        c.font, c.fill = head, fill
        c.alignment = Alignment(vertical="center")
        ws.column_dimensions[get_column_letter(i)].width = width

    rows = tl.sorted_entries(layers or _all_layers(), thread=thread)
    italic = Font(italic=True, color="6B6B6B")
    for r, e in enumerate(rows, start=2):
        for i, (_label, attr, _w, _in) in enumerate(COLUMNS, start=1):
            value = _as_date(e.when) if attr == "when" else _cell(e, attr)
            if attr == "doc_number" and e.doc_number:
                value = e.doc_number
            c = ws.cell(row=r, column=i, value=value)
            if attr == "when" and isinstance(value, datetime):
                c.number_format = "yyyy-mm-dd hh:mm"
            if attr == "title":
                c.alignment = Alignment(wrap_text=True, vertical="top")
            # An inference must never look like a record on the page.
            if e.layer == DERIVED:
                c.font = italic

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(COLUMNS))}{max(2, len(rows) + 1)}"

    about = wb.create_sheet("About this file")
    about.column_dimensions["A"].width = 110
    about["A1"] = f"Case {tl.case_number} — reconstructed timeline"
    about["A1"].font = Font(bold=True, size=13)
    about["A3"] = DISCLAIMER
    about["A3"].alignment = Alignment(wrap_text=True, vertical="top")
    about.row_dimensions[3].height = 60
    row = 5
    about[f"A{row}"] = "Layers"
    about[f"A{row}"].font = Font(bold=True)
    for text in (
        "Record — served, filed, or a court event. Of record.",
        "Correspondence — threads with counsel or third parties. Context, not record.",
        "Client — communications with the client. Privileged.",
        "Derived — our inference (a gap, a computed date). NOT a record; shown in italics.",
    ):
        row += 1
        about[f"A{row}"] = text
    if tl.gaps:
        row += 2
        about[f"A{row}"] = "Gaps"
        about[f"A{row}"].font = Font(bold=True)
        for g in tl.gaps:
            row += 1
            about[f"A{row}"] = f"[{g['class']}] {g['detail']}"
            about[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")
    if tl.findings:
        row += 2
        about[f"A{row}"] = "Findings"
        about[f"A{row}"].font = Font(bold=True)
        for f in tl.findings:
            row += 1
            about[f"A{row}"] = f
            about[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(path)
    return path


def _all_layers():
    from .timeline import LAYERS
    return LAYERS


def to_docx(tl: Timeline, path: str | Path, *, layers=None, thread=None) -> Path:
    """A real Word table, landscape, header repeating on every page."""
    try:
        import docx
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        raise RuntimeError(
            "Word export needs the 'docx' extra: pip install 'docketry[docx]'"
        ) from None

    path = Path(path)
    doc = docx.Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    # Narrow margins: a nine-column timeline needs the width more than it
    # needs a generous gutter.
    sec.left_margin = sec.right_margin = Inches(0.5)
    sec.top_margin = sec.bottom_margin = Inches(0.55)

    doc.add_heading(f"Case {tl.case_number} — Timeline", level=1)
    note = doc.add_paragraph(DISCLAIMER)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in note.runs:
        run.italic = True
        run.font.size = Pt(8.5)

    rows = tl.sorted_entries(layers or _all_layers(), thread=thread)
    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Light Grid Accent 1"
    # Three things have to agree or Word quietly autofits and hands every
    # column the same width: layout fixed, the table GRID, and each cell.
    # Setting only cell widths (the obvious call) does nothing at all.
    table.autofit = False
    table.allow_autofit = False
    for i, col in enumerate(table.columns):
        col.width = Inches(COLUMNS[i][3])
    for i, (label, _a, _w, _in) in enumerate(COLUMNS):
        cell = table.rows[0].cells[i]
        cell.text = label
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
        cell.width = Inches(COLUMNS[i][3])
    # Repeat the header when the table breaks across pages — without this a
    # printed timeline's later pages are unlabelled columns.
    table.rows[0]._tr.get_or_add_trPr().append(
        docx.oxml.parse_xml('<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/'
                            'wordprocessingml/2006/main" w:val="true"/>')
    )

    for e in rows:
        cells = table.add_row().cells
        for i, (_label, attr, _w, _in) in enumerate(COLUMNS):
            text = _cell(e, attr)
            if attr == "when" and e.when:
                text = e.when[:16].replace("T", " ")
            cells[i].text = text
            cells[i].width = Inches(COLUMNS[i][3])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    # Addresses have no spaces to break at, so they get a
                    # smaller face rather than being sliced mid-domain.
                    r.font.size = Pt(7) if attr == "actor" else Pt(8.5)
                    if e.layer == DERIVED:
                        r.italic = True

    if tl.gaps or tl.findings:
        doc.add_page_break()
        doc.add_heading("Gaps and findings", level=2)
        for g in tl.gaps:
            doc.add_paragraph(f"[{g['class']}] {g['detail']}", style="List Bullet")
        for f in tl.findings:
            doc.add_paragraph(f, style="List Bullet")

    doc.save(path)
    return path
